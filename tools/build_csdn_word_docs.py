from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "csdn_word_docs"
IMAGE_DIR = OUTPUT_DIR / "images"

FONT_REGULAR = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_BOLD = Path(r"C:\Windows\Fonts\msyhbd.ttc")


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = FONT_BOLD if bold and FONT_BOLD.exists() else FONT_REGULAR
    return ImageFont.truetype(str(path), size=size)


PALETTE = {
    "navy": "#1D3557",
    "blue": "#4C8BF5",
    "sky": "#A8DADC",
    "mint": "#CDEFE6",
    "sand": "#F6E7CB",
    "rose": "#F9D7D4",
    "yellow": "#FFD166",
    "green": "#4CAF50",
    "bg": "#FFF8F1",
    "ink": "#2F3441",
    "muted": "#64748B",
    "line": "#D7DFEA",
}


def rounded_box(draw: ImageDraw.ImageDraw, box, fill, outline=None, radius=28, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw: ImageDraw.ImageDraw, start, end, color, width=6, head=18):
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        sign = 1 if ex >= sx else -1
        draw.polygon(
            [(ex, ey), (ex - sign * head, ey - head // 2), (ex - sign * head, ey + head // 2)],
            fill=color,
        )
    else:
        sign = 1 if ey >= sy else -1
        draw.polygon(
            [(ex, ey), (ex - head // 2, ey - sign * head), (ex + head // 2, ey - sign * head)],
            fill=color,
        )


def draw_character(draw: ImageDraw.ImageDraw, x: int, y: int, body: str, accent: str):
    draw.ellipse((x, y, x + 70, y + 70), fill="#FFE0BD", outline=PALETTE["ink"], width=3)
    draw.ellipse((x + 20, y + 22, x + 28, y + 30), fill=PALETTE["ink"])
    draw.ellipse((x + 42, y + 22, x + 50, y + 30), fill=PALETTE["ink"])
    draw.arc((x + 18, y + 30, x + 52, y + 50), start=10, end=170, fill=PALETTE["ink"], width=3)
    rounded_box(draw, (x - 2, y + 70, x + 72, y + 160), fill=body, outline=PALETTE["ink"], radius=22)
    draw.line((x + 18, y + 160, x + 12, y + 208), fill=PALETTE["ink"], width=5)
    draw.line((x + 52, y + 160, x + 58, y + 208), fill=PALETTE["ink"], width=5)
    draw.line((x - 4, y + 102, x - 28, y + 140), fill=PALETTE["ink"], width=5)
    draw.line((x + 74, y + 102, x + 102, y + 126), fill=PALETTE["ink"], width=5)
    draw.ellipse((x + 84, y + 114, x + 104, y + 134), fill=accent, outline=PALETTE["ink"], width=3)


def write_center(draw: ImageDraw.ImageDraw, text: str, font, xy, fill):
    bbox = draw.textbbox((0, 0), text, font=font)
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    x, y = xy
    draw.text((x - w / 2, y - h / 2), text, font=font, fill=fill)


def create_cover(article_no: int, title: str, subtitle: str, theme: str) -> Path:
    image = Image.new("RGB", (1600, 900), PALETTE["bg"])
    draw = ImageDraw.Draw(image)

    rounded_box(draw, (70, 60, 1530, 840), fill="#FFFFFF", outline=PALETTE["line"], radius=40, width=4)
    rounded_box(draw, (110, 110, 770, 770), fill=PALETTE["sand"], radius=36)
    rounded_box(draw, (810, 110, 1490, 770), fill="#F8FBFF", radius=36)
    rounded_box(draw, (120, 120, 360, 170), fill=PALETTE["navy"], radius=20)
    draw.text((150, 131), f"第 {article_no} 篇", font=load_font(34, bold=True), fill="#FFFFFF")

    draw.text((120, 220), title, font=load_font(54, bold=True), fill=PALETTE["ink"])
    draw.text((120, 310), subtitle, font=load_font(28), fill=PALETTE["muted"])

    draw_character(draw, 180, 470, PALETTE["blue"], PALETTE["yellow"])
    draw_character(draw, 330, 430, PALETTE["mint"], PALETTE["rose"])
    draw_character(draw, 495, 500, PALETTE["rose"], PALETTE["sky"])

    rounded_box(draw, (855, 170, 1440, 300), fill=PALETTE["blue"], radius=28)
    draw.text((895, 205), theme, font=load_font(36, bold=True), fill="#FFFFFF")
    draw.text((895, 255), "软件工程视角", font=load_font(26), fill="#EAF4FF")

    for idx, (label, fill, y) in enumerate(
        [
            ("需求", PALETTE["sand"], 360),
            ("架构", PALETTE["mint"], 470),
            ("开发", PALETTE["sky"], 580),
        ]
    ):
        rounded_box(draw, (900, y, 1110, y + 76), fill=fill, outline=PALETTE["line"], radius=22)
        draw.text((958, y + 22), label, font=load_font(28, bold=True), fill=PALETTE["ink"])
        rounded_box(draw, (1170, y, 1380, y + 76), fill="#FFFFFF", outline=PALETTE["line"], radius=22)
        draw.text((1228, y + 22), "落地", font=load_font(28, bold=True), fill=PALETTE["blue"])
        arrow(draw, (1116, y + 38), (1164, y + 38), PALETTE["blue"], width=5, head=16)

    out = IMAGE_DIR / f"{article_no:02d}_cover.png"
    image.save(out)
    return out


def create_flow_image(article_no: int, title: str, steps: list[str], accent: str) -> Path:
    image = Image.new("RGB", (1600, 900), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, 1600, 900), fill="#FFFFFF")
    draw.text((80, 60), title, font=load_font(44, bold=True), fill=PALETTE["ink"])
    draw.text((80, 120), "用一张图快速看懂这一篇的核心逻辑", font=load_font(24), fill=PALETTE["muted"])

    xs = [110, 400, 690, 980, 1270]
    y = 340
    colors = [PALETTE["sand"], PALETTE["mint"], PALETTE["sky"], PALETTE["rose"], "#EAE4FF"]
    for idx, step in enumerate(steps):
        x = xs[idx]
        rounded_box(draw, (x, y, x + 210, y + 150), fill=colors[idx], outline=PALETTE["line"], radius=28)
        rounded_box(draw, (x + 18, y - 34, x + 84, y + 22), fill=accent, radius=18)
        write_center(draw, str(idx + 1), load_font(28, bold=True), (x + 51, y - 7), "#FFFFFF")

        lines = step.split("\n")
        ty = y + 44
        for line in lines:
            write_center(draw, line, load_font(28, bold=True), (x + 105, ty), PALETTE["ink"])
            ty += 38

        if idx < len(steps) - 1:
            arrow(draw, (x + 214, y + 75), (xs[idx + 1] - 14, y + 75), accent, width=6, head=18)

    rounded_box(draw, (120, 620, 1480, 790), fill="#F8FBFF", outline=PALETTE["line"], radius=32)
    draw.text((160, 660), "写作提示：每一小节都围绕“为什么这样做、在项目里怎么做、这样做带来什么效果”展开。", font=load_font(28), fill=PALETTE["ink"])

    out = IMAGE_DIR / f"{article_no:02d}_flow.png"
    image.save(out)
    return out


def create_repo_map() -> Path:
    image = Image.new("RGB", (1600, 1000), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    draw.text((80, 60), "仓库模块拆分图", font=load_font(44, bold=True), fill=PALETTE["ink"])

    center = (800, 180)
    rounded_box(draw, (620, 120, 980, 240), fill=PALETTE["navy"], radius=28)
    write_center(draw, "人工智能实践", load_font(40, bold=True), center, "#FFFFFF")

    nodes = [
        ("backend", (160, 390), PALETTE["sand"], "统一业务和接口"),
        ("pc-web", (470, 390), PALETTE["mint"], "用户端 PC 展示"),
        ("admin-web", (780, 390), PALETTE["sky"], "运营管理后台"),
        ("miniapp", (1090, 390), PALETTE["rose"], "微信小程序"),
        ("docs", (640, 690), "#EAE4FF", "需求/架构/验收文档"),
    ]
    for name, (x, y), fill, desc in nodes:
        rounded_box(draw, (x, y, x + 250, y + 140), fill=fill, outline=PALETTE["line"], radius=28)
        draw.text((x + 26, y + 28), name, font=load_font(34, bold=True), fill=PALETTE["ink"])
        draw.text((x + 26, y + 82), desc, font=load_font(22), fill=PALETTE["muted"])
        arrow(draw, (800, 240), (x + 125, y - 18), PALETTE["blue"], width=5, head=16)

    out = IMAGE_DIR / "05_repo_map.png"
    image.save(out)
    return out


def create_collage(name: str, title: str, image_paths: list[Path], labels: list[str]) -> Path:
    canvas = Image.new("RGB", (1600, 980), "#FFFFFF")
    draw = ImageDraw.Draw(canvas)
    draw.text((70, 50), title, font=load_font(42, bold=True), fill=PALETTE["ink"])
    draw.text((70, 110), "项目真实页面截图，可以直接作为文章里的案例图。", font=load_font(24), fill=PALETTE["muted"])

    slots = [(70, 190, 760, 610), (840, 190, 1530, 610), (420, 650, 1180, 930)]
    for idx, (img_path, label) in enumerate(zip(image_paths, labels)):
        src = Image.open(img_path).convert("RGB")
        box = slots[idx]
        x1, y1, x2, y2 = box
        area_w = x2 - x1
        area_h = y2 - y1
        ratio = min(area_w / src.width, area_h / src.height)
        resized = src.resize((int(src.width * ratio), int(src.height * ratio)))
        rounded_box(draw, box, fill="#F8FBFF", outline=PALETTE["line"], radius=24, width=3)
        rx = x1 + (area_w - resized.width) // 2
        ry = y1 + (area_h - resized.height) // 2
        canvas.paste(resized, (rx, ry))
        rounded_box(draw, (x1 + 18, y1 + 18, x1 + 220, y1 + 64), fill=PALETTE["navy"], radius=16)
        draw.text((x1 + 40, y1 + 28), label, font=load_font(22, bold=True), fill="#FFFFFF")

    out = IMAGE_DIR / name
    canvas.save(out)
    return out


def create_checklist_visual() -> Path:
    image = Image.new("RGB", (1600, 900), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    draw.text((80, 60), "联调与验收清单", font=load_font(44, bold=True), fill=PALETTE["ink"])
    draw.text((80, 120), "按这个顺序检查，课程项目也能做出正式交付的感觉。", font=load_font(24), fill=PALETTE["muted"])
    sections = [
        ("环境启动", ["后端能启动", "PC 端能打开", "后台能登录"]),
        ("接口检查", ["分类接口可返回", "详情接口字段完整", "收藏接口可用"]),
        ("页面检查", ["列表可浏览", "详情图文正常", "收藏状态同步"]),
        ("交付整理", ["SQL 脚本齐全", "说明文档完整", "截图和演示材料齐全"]),
    ]
    x = 100
    for title, items in sections:
        rounded_box(draw, (x, 220, x + 320, 700), fill="#F8FBFF", outline=PALETTE["line"], radius=28)
        rounded_box(draw, (x + 24, 244, x + 296, 314), fill=PALETTE["blue"], radius=18)
        write_center(draw, title, load_font(28, bold=True), (x + 160, 279), "#FFFFFF")
        iy = 370
        for item in items:
            draw.ellipse((x + 34, iy - 6, x + 66, iy + 26), fill=PALETTE["green"])
            draw.line((x + 43, iy + 9, x + 50, iy + 18), fill="#FFFFFF", width=4)
            draw.line((x + 50, iy + 18, x + 60, iy - 2), fill="#FFFFFF", width=4)
            draw.text((x + 82, iy - 10), item, font=load_font(24), fill=PALETTE["ink"])
            iy += 86
        x += 360
    out = IMAGE_DIR / "06_checklist.png"
    image.save(out)
    return out


def set_cell_text(cell, text: str):
    cell.text = text
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            set_run_font(run, 10.5)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, size: float, bold: bool = False, color: str = PALETTE["ink"], name: str = "Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))


def style_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, PALETTE["blue"], 18, 10),
        ("Heading 2", 13, PALETTE["blue"], 14, 7),
        ("Heading 3", 12, PALETTE["navy"], 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    if "CodeBlock" not in styles:
        code_style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Consolas"
        code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        code_style.font.size = Pt(10.5)
        code_style.paragraph_format.left_indent = Cm(0.5)
        code_style.paragraph_format.right_indent = Cm(0.5)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        code_style.paragraph_format.line_spacing = 1.15


def add_title(doc: Document, title: str, subtitle: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title)
    set_run_font(run, 22, bold=True, color=PALETTE["navy"])

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_after = Pt(12)
    run = sp.add_run(subtitle)
    set_run_font(run, 11, color=PALETTE["muted"])


def add_meta_bar(doc: Document, article_no: int):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [Inches(1.2), Inches(2.3), Inches(3.0)]
    labels = [
        f"系列第 {article_no} 篇",
        "主题：软件工程视角",
        "定位：可直接复制到 CSDN 的简洁实战文",
    ]
    for idx, (cell, width, text) in enumerate(zip(table.rows[0].cells, widths, labels)):
        cell.width = width
        set_cell_text(cell, text)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F8FBFF")
        tc_pr.append(shd)
    doc.add_paragraph()


def add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    set_run_font(run, 11)
    return p


def add_bullets(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run, 11)


def add_numbered(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run, 11)


def add_code_block(doc: Document, lines: list[str]):
    for line in lines:
        p = doc.add_paragraph(style="CodeBlock")
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.right_indent = Cm(0.7)
        p.paragraph_format.line_spacing = 1.15
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F4F6F9")
        pPr.append(shd)
        run = p.add_run(line)
        set_run_font(run, 10.5, name="Consolas")
    doc.add_paragraph()


def add_image(doc: Document, path: Path, caption: str, width: float = 6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(8)
    run = cp.add_run(caption)
    set_run_font(run, 10.5, color=PALETTE["muted"])


@dataclass
class Section:
    heading: str
    paragraphs: list[str] = field(default_factory=list)
    bullets: list[str] = field(default_factory=list)
    numbered: list[str] = field(default_factory=list)
    code: list[str] = field(default_factory=list)


@dataclass
class Article:
    no: int
    title: str
    subtitle: str
    theme: str
    flow_title: str
    flow_steps: list[str]
    sections: list[Section]
    extra_images: list[tuple[Path, str, float]] = field(default_factory=list)


ARTICLES = [
    Article(
        no=1,
        title="不是先写代码，而是先搭框架",
        subtitle="用软件工程思想启动一个四端项目",
        theme="项目启动与框架搭建",
        flow_title="项目启动的正确顺序",
        flow_steps=["先想目标", "再拆需求", "再定架构", "先跑框架", "最后开写"],
        sections=[
            Section(
                heading="引子",
                paragraphs=[
                    "很多同学做课程项目时，第一反应就是建表、写接口、画页面。结果常常是代码写了一半，需求改了；页面写出来了，接口却对不上；项目能运行，但越做越乱。",
                    "真正能把项目做顺的人，第一步通常不是写代码，而是先搭框架。这里说的框架，不只是代码框架，更是需求框架、模块框架和开发顺序。",
                ],
            ),
            Section(
                heading="1. 为什么不能一上来就写代码",
                bullets=[
                    "需求没定清楚，后面一定返工。",
                    "模块没拆清楚，团队协作会互相卡住。",
                    "接口没统一，联调阶段会特别痛苦。",
                    "验收标准没提前想好，最后只能临时补材料。",
                ],
                paragraphs=[
                    "软件工程最重要的价值，不是把流程搞复杂，而是帮我们把风险前移。把该想清楚的事情提前想清楚，后面的开发速度反而会更快。"
                ],
            ),
            Section(
                heading="2. 先把项目目标说清楚",
                paragraphs=[
                    "我这个项目的目标不是做一个大而全的平台，而是做一个能够演示完整业务链路的非遗文旅商品系统。它至少要覆盖商品浏览、商品详情、收藏、后台管理和多端联调。",
                    "目标一旦明确，后面的取舍就会简单很多。比如订单、支付、物流这些功能虽然常见，但不属于这次最小可交付范围，就可以暂时不做。"
                ],
            ),
            Section(
                heading="3. 再把系统拆成四端",
                bullets=[
                    "backend：统一业务规则、数据访问和接口输出。",
                    "pc-web：面向普通用户的 PC 端浏览体验。",
                    "admin-web：面向管理员的分类、商品和上下架管理。",
                    "miniapp：面向移动端场景的小程序浏览与收藏。",
                ],
                paragraphs=[
                    "这样拆的好处是职责清楚。每个端只关心自己的表现层，真正的业务规则收口在后端，后面联调也更稳定。"
                ],
            ),
            Section(
                heading="4. 搭一个最小可运行版本",
                numbered=[
                    "先让后端能启动，并提供最基本的商品与分类接口。",
                    "再让 PC 端和后台端先把页面骨架跑起来。",
                    "然后接上数据库和真实接口，替换静态数据。",
                    "最后再补小程序端和联调检查。",
                ],
                paragraphs=[
                    "最小可运行版本的意义，是先把主链路打通。只要主链路通了，后面再补细节，心里就不会慌。"
                ],
                code=[
                    "cd backend",
                    "npm run build",
                    "npm run start:local",
                    "",
                    "cd pc-web",
                    "npm run dev",
                    "",
                    "cd admin-web",
                    "npm run dev",
                ],
            ),
            Section(
                heading="5. 这个项目里的真实落地方式",
                paragraphs=[
                    "从当前仓库来看，项目已经按软件工程的思路做了分层：`backend` 负责接口，`pc-web` 和 `miniapp` 负责用户端体验，`admin-web` 负责运营管理，`docs` 负责需求、架构、数据库和验收说明。",
                    "这说明项目不是从某一个页面开始堆出来的，而是按“先设计，再实现，再联调”的节奏推进的。"
                ],
            ),
            Section(
                heading="结语",
                paragraphs=[
                    "项目真正的起点，不是第一行代码，而是第一张清晰的框架图。只要前面的目标、模块和顺序定清楚，后面的开发就不会乱。"
                ],
            ),
        ],
    ),
    Article(
        no=2,
        title="需求分析怎么做",
        subtitle="先把角色、功能和边界讲清楚",
        theme="角色与需求边界",
        flow_title="需求分析的核心动作",
        flow_steps=["找用户", "定角色", "列功能", "划边界", "出清单"],
        sections=[
            Section(
                heading="引子",
                paragraphs=[
                    "很多项目不是死在技术上，而是死在需求模糊上。大家口头上都知道要做一个系统，但到底给谁用、做哪些功能、哪些功能不做，往往没有提前说清楚。",
                    "需求分析不是写一大堆空话，而是把项目的边界画出来。边界越清楚，后面的开发越省力。"
                ],
            ),
            Section(
                heading="1. 先找出项目里的核心角色",
                bullets=[
                    "游客：不登录也能看商品列表和详情。",
                    "普通用户：登录后可以收藏、取消收藏、查看我的收藏。",
                    "管理员：登录后台后可以管理分类、商品、SKU、图片和上下架状态。",
                ],
                paragraphs=[
                    "角色一旦分清楚，功能权限就自然有了边界。谁能看，谁能改，谁能操作后台，都能顺着角色往下拆。"
                ],
            ),
            Section(
                heading="2. 再把功能模块拆出来",
                bullets=[
                    "商品分类模块",
                    "商品浏览模块",
                    "商品详情模块",
                    "收藏模块",
                    "后台商品管理模块",
                    "后台分类和状态管理模块",
                ],
                paragraphs=[
                    "这里的关键不是写得越多越好，而是写出最小闭环。比如有了商品浏览，就最好有商品详情；有了收藏，就最好有收藏列表。"
                ],
            ),
            Section(
                heading="3. 明确哪些属于本期范围",
                paragraphs=[
                    "当前项目聚焦的是非遗商品展示与管理，所以重点放在商品内容本身，而不是完整交易链路。订单、支付、物流这些模块都可以明确标成后续扩展。",
                    "这样做并不是偷懒，而是典型的软件工程思维：先完成最小可交付版本，再谈更大的系统能力。"
                ],
            ),
            Section(
                heading="4. 为什么边界比功能数量更重要",
                bullets=[
                    "边界清楚，开发任务更容易拆分。",
                    "边界清楚，接口设计不会反复推翻。",
                    "边界清楚，验收时就知道该检查什么。",
                ],
                paragraphs=[
                    "课程项目最怕的不是功能少，而是做了一堆半成品。与其把系统做得很大，不如把关键链路做完整。"
                ],
            ),
            Section(
                heading="5. 本项目适合输出哪些需求文档",
                bullets=[
                    "角色说明：谁在用系统。",
                    "功能清单：系统至少要完成什么。",
                    "权限边界：哪些接口是公开的，哪些需要登录。",
                    "不做事项：本期明确不覆盖什么内容。",
                ],
            ),
            Section(
                heading="结语",
                paragraphs=[
                    "需求分析不是为了写文档而写文档，而是为了让团队知道自己到底在做什么。角色、功能和边界一旦清楚，后面的架构设计才会稳。"
                ],
            ),
        ],
    ),
    Article(
        no=3,
        title="系统架构设计怎么做",
        subtitle="为什么这个项目适合做成四端联动",
        theme="四端架构与统一接口",
        flow_title="架构设计的主线",
        flow_steps=["定职责", "定接口", "定数据流", "定端口", "定联调"],
        sections=[
            Section(
                heading="引子",
                paragraphs=[
                    "当一个项目同时有 PC 端、管理后台和小程序时，如果没有架构设计，最后很容易变成多套逻辑并行、接口命名混乱、前后端重复造轮子。",
                    "这类项目的关键不是端越多越厉害，而是多端之间能不能共用一套稳定的后端能力。"
                ],
            ),
            Section(
                heading="1. 为什么拆成四端",
                bullets=[
                    "MySQL：负责持久化分类、商品、SKU、图片和收藏数据。",
                    "backend：负责统一的业务规则、鉴权和接口输出。",
                    "pc-web：负责 PC 场景下的浏览展示。",
                    "admin-web：负责运营和商品维护。",
                    "miniapp：负责移动端浏览和收藏。",
                ],
                paragraphs=[
                    "四端并不是四套系统，而是一套后端支撑多种使用场景。真正的重点是职责分离，而不是简单地把目录拆开。"
                ],
            ),
            Section(
                heading="2. 为什么三端前端共用一套接口",
                paragraphs=[
                    "如果 PC 端一套接口，小程序一套接口，后台再来一套接口，后续维护成本会很高。相同的数据口径会被复制三次，业务规则也容易不一致。",
                    "更好的方式是让商品、分类、收藏这些规则统一沉到后端。前端只负责展示，接口尽量共用，必要时再按角色区分权限。"
                ],
            ),
            Section(
                heading="3. 架构图里最重要的不是图，而是流转关系",
                numbered=[
                    "用户在 PC 端或小程序端发起浏览请求。",
                    "请求进入统一后端接口层。",
                    "后端根据上架状态、删除状态和鉴权规则过滤数据。",
                    "最后再把结果返回给不同前端。",
                ],
                paragraphs=[
                    "这样一来，商品是否可见、收藏是否有效，不需要每个前端自己判断，而是由后端统一控制。"
                ],
            ),
            Section(
                heading="4. 本项目里的实际运行端口",
                code=[
                    "backend  -> http://127.0.0.1:3000/api",
                    "pc-web   -> http://127.0.0.1:5173",
                    "admin-web-> http://127.0.0.1:5174",
                    "miniapp  -> 微信开发者工具打开 miniapp 目录",
                ],
                paragraphs=[
                    "把本地端口和访问入口统一记录下来，联调时会轻松很多。否则到了验收阶段，大家连项目从哪里打开都要重新确认。"
                ],
            ),
            Section(
                heading="5. 一个好架构会带来什么",
                bullets=[
                    "新增一个前端端口时，不需要重写一套业务逻辑。",
                    "管理端修改商品状态后，用户端能立即感知变化。",
                    "联调时问题定位更快，因为责任边界很清楚。",
                ],
            ),
            Section(
                heading="结语",
                paragraphs=[
                    "架构设计的目的，不是把图画复杂，而是让每个模块只做自己该做的事。这个项目之所以适合做成四端联动，核心就在于一套后端能力服务多端场景。"
                ],
            ),
        ],
    ),
    Article(
        no=4,
        title="数据库和接口设计怎么落地",
        subtitle="先把数据想清楚，后面才不会乱",
        theme="数据模型与接口分层",
        flow_title="从数据到接口的设计顺序",
        flow_steps=["找对象", "连关系", "定字段", "分接口", "统一返回"],
        sections=[
            Section(
                heading="引子",
                paragraphs=[
                    "很多同学做项目时，会把数据库设计和接口设计分开看。其实这两件事应该连起来考虑，因为接口只是数据模型在业务层的一种输出方式。",
                    "如果数据关系没想清楚，后面的接口一定会越来越别扭。"
                ],
            ),
            Section(
                heading="1. 先找出项目里的核心数据对象",
                bullets=[
                    "product_category：商品分类",
                    "product：商品主体",
                    "product_sku：规格、价格、库存",
                    "product_image：商品图片",
                    "favorite：用户收藏",
                ],
                paragraphs=[
                    "这些对象不是随便列的，而是从业务里抽出来的。用户能看到什么、管理员能维护什么，背后都能映射到这些实体。"
                ],
            ),
            Section(
                heading="2. 再看对象之间的关系",
                bullets=[
                    "一个分类下面可以有多个商品。",
                    "一个商品下面可以有多个 SKU。",
                    "一个商品下面可以有多张图片。",
                    "一个用户可以收藏多个商品。",
                ],
                paragraphs=[
                    "把这些关系先确定下来，后面的接口路径就更自然。比如详情页为什么能同时返回商品、图片和 SKU，本质上就是因为数据关系提前设计好了。"
                ],
            ),
            Section(
                heading="3. 接口不要一股脑混在一起",
                paragraphs=[
                    "这个项目里最合理的做法，是把公开接口和管理接口分开。公开接口面向游客和普通用户，管理接口面向后台管理员。",
                    "这样一来，用户端只能拿到可展示的数据，而管理端能看到更完整的维护信息。权限边界和接口边界也就同步清楚了。"
                ],
                bullets=[
                    "公开接口：分类列表、商品分页、商品详情、收藏相关接口。",
                    "管理接口：分类 CRUD、商品 CRUD、SKU 管理、图片管理、上下架管理。",
                ],
            ),
            Section(
                heading="4. 返回结构为什么要统一",
                bullets=[
                    "前端更容易封装请求工具。",
                    "联调时更容易判断问题出在字段还是业务。",
                    "后续接口扩展不会把页面层全部推翻。",
                ],
                paragraphs=[
                    "一个项目里真正可怕的，不是字段多，而是同一类接口风格完全不一致。统一返回结构，是后期维护体验的基础。"
                ],
            ),
            Section(
                heading="5. 先设计数据，再写接口的好处",
                paragraphs=[
                    "你会发现，一旦分类、商品、SKU、图片和收藏这条主线建好了，前台浏览、后台管理和联调验收都会顺很多。因为它们看起来是不同功能，本质上用的是同一批核心数据。"
                ],
            ),
            Section(
                heading="结语",
                paragraphs=[
                    "数据库设计解决的是系统怎么存，接口设计解决的是系统怎么用。把这两件事连起来看，项目才会越做越稳。"
                ],
            ),
        ],
    ),
    Article(
        no=5,
        title="前后端分模块开发",
        subtitle="四个子项目是怎么一步步搭起来的",
        theme="模块拆分与并行开发",
        flow_title="四端并行开发的基本方法",
        flow_steps=["先分仓", "后端先行", "页面骨架", "对接接口", "统一验收"],
        sections=[
            Section(
                heading="引子",
                paragraphs=[
                    "一个四端项目最怕的事情，不是任务多，而是大家同时开工却没有边界。结果往往是后端不知道前端要什么，前端也不知道接口什么时候能稳定。",
                    "分模块开发的核心目标，就是让每个子项目都能独立推进，但又能在关键节点顺利汇合。"
                ],
            ),
            Section(
                heading="1. 先把仓库结构拆出来",
                bullets=[
                    "`backend`：Cool Admin Midway + Midway.js + MySQL",
                    "`pc-web`：React + Vite + React Router",
                    "`admin-web`：React + Vite + Ant Design",
                    "`miniapp`：原生微信小程序",
                    "`docs`：需求、架构、数据库、接口与验收文档",
                ],
                paragraphs=[
                    "这类目录结构最大的好处，就是一眼能看出每个模块在做什么。后面不管是自己开发还是给同学交接，理解成本都会低很多。"
                ],
            ),
            Section(
                heading="2. 后端模块为什么要先打底",
                paragraphs=[
                    "后端是整个项目的中枢。商品分类、商品详情、收藏逻辑、上下架规则，都应该在后端先稳定下来。",
                    "当前项目里，业务模块主要落在 `backend/src/modules/wudong` 下，说明它已经不是一个空框架，而是按业务模块继续扩展出来的。"
                ],
            ),
            Section(
                heading="3. 用户端页面怎么分工",
                bullets=[
                    "PC Web 负责列表、详情和收藏等浏览型页面。",
                    "小程序负责移动端同类场景，但更贴近轻量操作。",
                    "两端尽量共用接口，只在交互层面做差异化。",
                ],
            ),
            Section(
                heading="4. 管理后台怎么体现运营思维",
                paragraphs=[
                    "管理后台不是为了好看，而是为了让数据可维护。分类、商品、SKU、图片、上下架这些动作，决定了用户端最终能看到什么。",
                    "所以后台端的价值，不在页面多少，而在是否能让运营动作和用户结果形成闭环。"
                ],
            ),
            Section(
                heading="5. 并行开发时最关键的几件事",
                bullets=[
                    "先定接口，再定页面真实数据结构。",
                    "先跑通骨架，再补复杂交互。",
                    "前端能用静态数据时先别堵在后端上。",
                    "接口一旦联通，就尽快替换成真实请求。",
                ],
            ),
            Section(
                heading="结语",
                paragraphs=[
                    "模块拆分不是为了把目录拆得更细，而是为了让开发有节奏。一个四端项目只要职责清楚、接口统一，就能实现真正意义上的并行推进。"
                ],
            ),
        ],
    ),
    Article(
        no=6,
        title="联调、验收与交付",
        subtitle="一个项目能跑起来，不代表它真的完成了",
        theme="联调检查与交接收尾",
        flow_title="项目收尾的推荐顺序",
        flow_steps=["先启动", "查接口", "看页面", "补文档", "做交付"],
        sections=[
            Section(
                heading="引子",
                paragraphs=[
                    "很多课程项目做到最后，代码是有的，页面也能打开，但一到演示和交接就开始慌。其实问题不在于做得少，而在于没有把联调、验收和交付当成正式工作来做。",
                    "项目能运行，只说明它开始有结果；项目能验收，才说明它真正闭环了。"
                ],
            ),
            Section(
                heading="1. 联调顺序要先后端、再前端、最后多端联动",
                numbered=[
                    "先确认后端能正常启动，核心接口有返回。",
                    "再确认 PC 端和后台端能打开，并接上真实接口。",
                    "最后再在小程序里检查浏览、详情和收藏链路。",
                ],
                paragraphs=[
                    "这样做的原因很简单：接口没有稳定之前，前端的很多问题其实是伪问题。先把后端打稳，联调才不会来回扯。"
                ],
            ),
            Section(
                heading="2. 联调阶段重点检查什么",
                bullets=[
                    "分类接口是否只返回可用分类。",
                    "商品列表和详情接口字段是否完整。",
                    "收藏状态是否能新增、取消和同步展示。",
                    "后台上下架后，用户端是否能立即反映变化。",
                ],
            ),
            Section(
                heading="3. 页面验收不要只看能不能打开",
                bullets=[
                    "列表是否能正常切换分类。",
                    "详情页的图文信息是否完整。",
                    "收藏页是否能看到真实数据变化。",
                    "后台页面是否能完成真正的管理动作。",
                ],
                paragraphs=[
                    "验收一定要从使用者视角去看，而不是只站在开发者视角看控制台有没有报错。"
                ],
            ),
            Section(
                heading="4. 交付材料为什么同样重要",
                bullets=[
                    "README 或运行说明，让别人知道怎么启动项目。",
                    "数据库脚本，让环境能快速恢复。",
                    "接口文档和页面截图，让演示有依据。",
                    "项目交接说明，让后续维护者快速接手。",
                ],
                paragraphs=[
                    "当前仓库里已经有 `docs`、SQL 脚本、交接说明和截图，这其实就是比较完整的交付意识。"
                ],
            ),
            Section(
                heading="5. 课程项目也可以做出正式感",
                paragraphs=[
                    "所谓正式感，并不是页面多华丽，而是别人接过你的项目时，能跑、能看、能理解、能继续改。只要把联调顺序、验收清单和交付材料做好，课程项目一样能很专业。"
                ],
            ),
            Section(
                heading="结语",
                paragraphs=[
                    "真正完成一个项目，最后拼的不是写代码的速度，而是把结果收束起来的能力。联调、验收和交付，决定了这个项目是不是一个完整作品。"
                ],
            ),
        ],
    ),
]


def build_doc(article: Article):
    doc = Document()
    style_document(doc)
    add_title(doc, article.title, article.subtitle)
    add_meta_bar(doc, article.no)

    cover = create_cover(article.no, article.title, article.subtitle, article.theme)
    flow = create_flow_image(article.no, article.flow_title, article.flow_steps, PALETTE["blue"])

    add_image(doc, cover, "图 1  卡通封面图：对应本篇主题", width=6.2)
    add_image(doc, flow, "图 2  逻辑流程图：适合直接插入 CSDN 作为主图", width=6.2)

    if article.no == 5:
        repo_map = create_repo_map()
        collage = create_collage(
            "05_screens.png",
            "项目真实页面案例图",
            [
                ROOT / "desktop-home-v4.png",
                ROOT / "detail-final-v2.png",
                ROOT / "favorites-mobile-final.png",
            ],
            ["PC 首页", "详情页", "小程序收藏页"],
        )
        article.extra_images.extend(
            [
                (repo_map, "图 3  仓库模块拆分图", 6.1),
                (collage, "图 4  项目真实页面截图拼图", 6.1),
            ]
        )

    if article.no == 6:
        checklist = create_checklist_visual()
        collage = create_collage(
            "06_acceptance.png",
            "联调与验收相关页面图",
            [
                ROOT / "desktop-home-v4.png",
                ROOT / "detail-final-v2.png",
                ROOT / "favorites-final-v2.png",
            ],
            ["列表页", "详情页", "收藏页"],
        )
        article.extra_images.extend(
            [
                (checklist, "图 3  联调与验收清单图", 6.1),
                (collage, "图 4  页面验收示例图", 6.1),
            ]
        )

    for section in article.sections:
        doc.add_heading(section.heading, level=1 if section.heading in {"引子", "结语"} else 2)
        for paragraph in section.paragraphs:
            add_body_paragraph(doc, paragraph)
        if section.bullets:
            add_bullets(doc, section.bullets)
        if section.numbered:
            add_numbered(doc, section.numbered)
        if section.code:
            add_code_block(doc, section.code)

    for path, caption, width in article.extra_images:
        add_image(doc, path, caption, width=width)

    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CSDN 系列文章文档版 | 适合复制后再按平台样式微调")
    set_run_font(run, 9.5, color=PALETTE["muted"])

    out = OUTPUT_DIR / f"{article.no:02d}_{article.title}.docx"
    doc.save(out)
    return out


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    paths = []
    for article in ARTICLES:
        paths.append(build_doc(article))

    index = ROOT / "output" / "csdn_word_docs" / "README.txt"
    lines = ["已生成的 Word 文档："]
    lines.extend(str(path.name) for path in paths)
    index.write_text("\n".join(lines), encoding="utf-8")


if __name__ == "__main__":
    main()
